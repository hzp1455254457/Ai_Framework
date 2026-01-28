"""
Resume模块 - 简历解析器

支持解析PDF、Word、JSON格式的简历文件，提取结构化数据。
"""

import re
import json
from typing import Dict, Any, Optional, List
from pathlib import Path
import io

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from docx import Document
except ImportError:
    Document = None

from core.base.service import BaseService
from core.resume.models import (
    ResumeData,
    PersonalInfo,
    Education,
    WorkExperience,
    ProjectExperience,
    Skill,
    Certificate,
)


class ResumeParseError(Exception):
    """简历解析错误"""
    pass


class ResumeParser(BaseService):
    """
    简历解析器
    
    支持解析PDF、Word、JSON格式的简历文件，提取结构化数据。
    
    支持的格式：
        - PDF: 使用pdfplumber提取文本和结构
        - Word: 使用python-docx提取内容
        - JSON: 直接加载和验证
    """
    
    def __init__(self, config: Dict[str, Any]) -> None:
        """
        初始化解析器
        
        参数:
            config: 配置字典，包含解析器相关配置
        """
        super().__init__(config)
        self.max_file_size = config.get("resume", {}).get("max_file_size", 10 * 1024 * 1024)  # 默认10MB
        self.supported_formats = ["pdf", "docx", "json"]
    
    async def parse(self, file_path: str, file_format: str) -> ResumeData:
        """
        解析简历文件
        
        参数:
            file_path: 文件路径
            file_format: 文件格式（pdf/docx/json）
        
        返回:
            ResumeData: 解析后的简历数据
        
        异常:
            ResumeParseError: 解析失败时抛出
        """
        if file_format.lower() not in self.supported_formats:
            raise ResumeParseError(f"不支持的文件格式: {file_format}")
        
        # 检查文件是否存在
        path = Path(file_path)
        if not path.exists():
            raise ResumeParseError(f"文件不存在: {file_path}")
        
        # 检查文件大小
        file_size = path.stat().st_size
        if file_size > self.max_file_size:
            raise ResumeParseError(f"文件大小超过限制: {file_size} bytes > {self.max_file_size} bytes")
        
        try:
            if file_format.lower() == "pdf":
                return await self._parse_pdf(file_path)
            elif file_format.lower() == "docx":
                return await self._parse_docx(file_path)
            elif file_format.lower() == "json":
                return await self._parse_json(file_path)
            else:
                raise ResumeParseError(f"不支持的文件格式: {file_format}")
        except Exception as e:
            self.logger.error(f"解析简历文件失败: {e}", exc_info=True)
            raise ResumeParseError(f"解析简历文件失败: {e}") from e
    
    async def _parse_pdf(self, file_path: str) -> ResumeData:
        """
        解析PDF文件
        
        参数:
            file_path: PDF文件路径
        
        返回:
            ResumeData: 解析后的简历数据
        """
        if pdfplumber is None:
            raise ResumeParseError("pdfplumber未安装，请先安装: pip install pdfplumber")
        
        try:
            # 使用pdfplumber提取文本
            text_content = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_content.append(text)
            
            full_text = "\n".join(text_content)
            if not full_text.strip():
                raise ResumeParseError("PDF文件未提取到文本内容")
            
            # 使用正则表达式提取结构化数据
            return self._extract_structured_data(full_text)
        except Exception as e:
            self.logger.error(f"解析PDF文件失败: {e}", exc_info=True)
            raise ResumeParseError(f"解析PDF文件失败: {e}") from e
    
    async def _parse_docx(self, file_path: str) -> ResumeData:
        """
        解析Word文件
        
        参数:
            file_path: Word文件路径
        
        返回:
            ResumeData: 解析后的简历数据
        """
        if Document is None:
            raise ResumeParseError("python-docx未安装，请先安装: pip install python-docx")
        
        try:
            # 使用python-docx提取内容
            doc = Document(file_path)
            text_content = []
            
            # 提取段落文本
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            full_text = "\n".join(text_content)
            if not full_text.strip():
                raise ResumeParseError("Word文件未提取到文本内容")
            
            # 使用正则表达式提取结构化数据
            return self._extract_structured_data(full_text)
        except Exception as e:
            self.logger.error(f"解析Word文件失败: {e}", exc_info=True)
            raise ResumeParseError(f"解析Word文件失败: {e}") from e
    
    async def _parse_json(self, file_path: str) -> ResumeData:
        """
        解析JSON文件
        
        参数:
            file_path: JSON文件路径
        
        返回:
            ResumeData: 解析后的简历数据
        """
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            
            # 验证并转换为ResumeData
            return ResumeData(**data)
        except json.JSONDecodeError as e:
            raise ResumeParseError(f"JSON格式错误: {e}")
        except Exception as e:
            self.logger.error(f"解析JSON文件失败: {e}", exc_info=True)
            raise ResumeParseError(f"解析JSON文件失败: {e}") from e
    
    def _extract_structured_data(self, text: str) -> ResumeData:
        """
        从文本中提取结构化数据
        
        使用正则表达式识别关键字段并构建ResumeData对象。
        
        参数:
            text: 简历文本内容
        
        返回:
            ResumeData: 结构化简历数据
        """
        # 提取个人信息
        personal_info = self._extract_personal_info(text)
        
        # 提取教育经历
        education = self._extract_education(text)
        
        # 提取工作经历
        work_experience = self._extract_work_experience(text)
        
        # 提取项目经历
        project_experience = self._extract_project_experience(text)
        
        # 提取技能
        skills = self._extract_skills(text)
        
        # 提取证书
        certificates = self._extract_certificates(text)
        
        return ResumeData(
            personal_info=personal_info,
            education=education,
            work_experience=work_experience,
            project_experience=project_experience,
            skills=skills,
            certificates=certificates,
        )
    
    def _extract_personal_info(self, text: str) -> PersonalInfo:
        """提取个人信息"""
        # 提取姓名（通常在开头）
        name_match = re.search(r"^[\u4e00-\u9fa5a-zA-Z\s]{2,20}", text.strip(), re.MULTILINE)
        name = name_match.group(0).strip() if name_match else "未知"
        
        # 提取邮箱
        email_match = re.search(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b", text)
        email = email_match.group(0) if email_match else ""
        
        # 提取电话
        phone_match = re.search(r"1[3-9]\d{9}|0\d{2,3}-?\d{7,8}", text)
        phone = phone_match.group(0) if phone_match else None
        
        # 提取地址
        location_match = re.search(r"(北京|上海|广州|深圳|杭州|成都|武汉|西安|南京|苏州|天津|重庆|长沙|郑州|青岛|大连|厦门|福州|济南|合肥|石家庄|哈尔滨|长春|沈阳|昆明|贵阳|南宁|海口|乌鲁木齐|拉萨|银川|西宁|呼和浩特)", text)
        location = location_match.group(0) if location_match else None
        
        # 提取网站链接
        website_match = re.search(r"https?://[^\s]+", text)
        website = website_match.group(0) if website_match else None
        
        # 提取LinkedIn
        linkedin_match = re.search(r"linkedin\.com/in/[^\s]+", text, re.IGNORECASE)
        linkedin = f"https://{linkedin_match.group(0)}" if linkedin_match else None
        
        # 提取GitHub
        github_match = re.search(r"github\.com/[^\s]+", text, re.IGNORECASE)
        github = f"https://{github_match.group(0)}" if github_match else None
        
        return PersonalInfo(
            name=name,
            email=email,
            phone=phone,
            location=location,
            website=website,
            linkedin=linkedin,
            github=github,
        )
    
    def _extract_education(self, text: str) -> List[Education]:
        """提取教育经历"""
        education_list = []
        
        # 查找教育部分（关键词：教育、学历、Education）
        edu_pattern = r"(教育|学历|Education|EDUCATION)[\s:：]*\n(.*?)(?=\n(工作|项目|技能|经验|Work|Project|Skill|Experience)|$)"
        edu_match = re.search(edu_pattern, text, re.DOTALL | re.IGNORECASE)
        
        if edu_match:
            edu_text = edu_match.group(2)
            # 提取学校、学位、专业、日期
            # 简化实现：查找学校名称和日期
            school_pattern = r"([\u4e00-\u9fa5a-zA-Z\s]+大学|[\u4e00-\u9fa5a-zA-Z\s]+学院|[\u4e00-\u9fa5a-zA-Z\s]+学校)"
            date_pattern = r"(\d{4}[-/]\d{1,2})[^\d]*(\d{4}[-/]\d{1,2})?"
            
            schools = re.findall(school_pattern, edu_text)
            dates = re.findall(date_pattern, edu_text)
            
            for i, school in enumerate(schools[:3]):  # 最多3个教育经历
                start_date = dates[i][0] if i < len(dates) else ""
                end_date = dates[i][1] if i < len(dates) and dates[i][1] else None
                
                education_list.append(Education(
                    school=school.strip(),
                    degree="本科",  # 默认值，实际应该从文本中提取
                    major="",  # 需要更复杂的提取逻辑
                    start_date=start_date,
                    end_date=end_date,
                ))
        
        return education_list
    
    def _extract_work_experience(self, text: str) -> List[WorkExperience]:
        """提取工作经历"""
        work_list = []
        
        # 查找工作部分（关键词：工作、经验、Experience、Work）
        work_pattern = r"(工作|经验|Experience|WORK)[\s:：]*\n(.*?)(?=\n(项目|技能|教育|Project|Skill|Education)|$)"
        work_match = re.search(work_pattern, text, re.DOTALL | re.IGNORECASE)
        
        if work_match:
            work_text = work_match.group(2)
            # 提取公司、职位、日期
            company_pattern = r"([\u4e00-\u9fa5a-zA-Z\s]+(公司|科技|集团|有限公司))"
            date_pattern = r"(\d{4}[-/]\d{1,2})[^\d]*(\d{4}[-/]\d{1,2})?"
            
            companies = re.findall(company_pattern, work_text)
            dates = re.findall(date_pattern, work_text)
            
            for i, company_match in enumerate(companies[:5]):  # 最多5个工作经历
                company = company_match[0] if isinstance(company_match, tuple) else company_match
                start_date = dates[i][0] if i < len(dates) else ""
                end_date = dates[i][1] if i < len(dates) and dates[i][1] else None
                
                work_list.append(WorkExperience(
                    company=company.strip(),
                    position="",  # 需要更复杂的提取逻辑
                    start_date=start_date,
                    end_date=end_date,
                    responsibilities=[],  # 需要更复杂的提取逻辑
                ))
        
        return work_list
    
    def _extract_project_experience(self, text: str) -> List[ProjectExperience]:
        """提取项目经历"""
        project_list = []
        
        # 查找项目部分（关键词：项目、Project）
        project_pattern = r"(项目|Project|PROJECT)[\s:：]*\n(.*?)(?=\n(技能|教育|工作|Skill|Education|Work)|$)"
        project_match = re.search(project_pattern, text, re.DOTALL | re.IGNORECASE)
        
        if project_match:
            project_text = project_match.group(2)
            # 简化实现：提取项目名称
            project_name_pattern = r"([\u4e00-\u9fa5a-zA-Z0-9\s]+项目|[\u4e00-\u9fa5a-zA-Z0-9\s]+系统)"
            projects = re.findall(project_name_pattern, project_text)
            
            for project_name in projects[:5]:  # 最多5个项目
                project_list.append(ProjectExperience(
                    name=project_name.strip(),
                    role="",
                    description="",
                    technologies=[],
                ))
        
        return project_list
    
    def _extract_skills(self, text: str) -> List[Skill]:
        """提取技能"""
        skills_list = []
        
        # 查找技能部分（关键词：技能、Skill）
        skill_pattern = r"(技能|Skill|SKILL)[\s:：]*\n(.*?)(?=\n(证书|教育|工作|Certificate|Education|Work)|$)"
        skill_match = re.search(skill_pattern, text, re.DOTALL | re.IGNORECASE)
        
        if skill_match:
            skill_text = skill_match.group(2)
            # 提取技能关键词（常见编程语言、框架、工具）
            common_skills = [
                "Python", "Java", "JavaScript", "TypeScript", "Go", "C++", "C#",
                "React", "Vue", "Angular", "Node.js", "Django", "Flask", "FastAPI",
                "MySQL", "PostgreSQL", "MongoDB", "Redis",
                "Docker", "Kubernetes", "AWS", "Azure", "Git",
            ]
            
            found_skills = []
            for skill in common_skills:
                if skill.lower() in skill_text.lower():
                    found_skills.append(skill)
            
            if found_skills:
                skills_list.append(Skill(
                    category="技术技能",
                    items=found_skills,
                ))
        
        return skills_list
    
    def _extract_certificates(self, text: str) -> List[Certificate]:
        """提取证书"""
        certificates_list = []
        
        # 查找证书部分（关键词：证书、认证、Certificate）
        cert_pattern = r"(证书|认证|Certificate|CERTIFICATE)[\s:：]*\n(.*?)(?=\n(技能|教育|工作|Skill|Education|Work)|$)"
        cert_match = re.search(cert_pattern, text, re.DOTALL | re.IGNORECASE)
        
        if cert_match:
            cert_text = cert_match.group(2)
            # 提取证书名称（简化实现）
            cert_name_pattern = r"([\u4e00-\u9fa5a-zA-Z\s]+(认证|证书|Certification))"
            certs = re.findall(cert_name_pattern, cert_text)
            
            for cert_match in certs[:5]:  # 最多5个证书
                cert_name = cert_match[0] if isinstance(cert_match, tuple) else cert_match
                certificates_list.append(Certificate(
                    name=cert_name.strip(),
                    issuer="",
                    date=None,
                ))
        
        return certificates_list
